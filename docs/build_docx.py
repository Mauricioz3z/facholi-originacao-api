# -*- coding: utf-8 -*-
"""Gera a documentação técnica da API em DOCX (identidade visual Facholi)."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- paleta ----------
GREEN_900 = "0E3A1B"; GREEN_800 = "134A22"; GREEN_700 = "195C27"
GREEN_600 = "1E6B30"; GREEN_500 = "23813C"; LEAF = "5C9E2E"
INK = "1C2620"; MUTED = "5C6B62"; LINE = "DFE6E0"; SOFT = "F4F7F4"
C_GET = "1E7D46"; C_POST = "1F6FB2"; C_PUT = "B9821A"; C_DEL = "B23B3B"
WHITE = "FFFFFF"

doc = Document()

# fonte padrão
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(4)
normal.paragraph_format.line_spacing = 1.15

def _shade(el, hexcolor):
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hexcolor)
    el.append(shd)

def cell_bg(cell, hexcolor):
    _shade(cell._tc.get_or_add_tcPr(), hexcolor)

def run_bg(run, hexcolor):
    _shade(run._r.get_or_add_rPr(), hexcolor)

def set_borders(tbl, color=LINE, sz=4, which="all"):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    edges = ["top","left","bottom","right"] + (["insideH","insideV"] if which=="all" else [])
    for edge in edges:
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"),"single"); e.set(qn("w:sz"),str(sz))
        e.set(qn("w:space"),"0"); e.set(qn("w:color"),color)
        borders.append(e)
    tblPr.append(borders)

def no_borders(tbl):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ["top","left","bottom","right","insideH","insideV"]:
        e = OxmlElement(f"w:{edge}"); e.set(qn("w:val"),"none"); borders.append(e)
    tblPr.append(borders)

def left_accent(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    e = OxmlElement("w:left"); e.set(qn("w:val"),"single")
    e.set(qn("w:sz"),"24"); e.set(qn("w:space"),"0"); e.set(qn("w:color"),color)
    borders.append(e); tcPr.append(borders)

def run(p, text, *, bold=False, color=INK, size=10.5, mono=False, italic=False):
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    r.font.name = "Consolas" if mono else "Calibri"
    return r

def para(text="", **kw):
    p = doc.add_paragraph()
    if text: run(p, text, **kw)
    return p

def heading(text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level==1 else 6)
    p.paragraph_format.space_after = Pt(3)
    if level == 1:
        run(p, text, bold=True, color=GREEN_700, size=13)
    else:
        run(p, text, bold=True, color=GREEN_800, size=11)
    return p

def section_header(num, title):
    doc.add_page_break()
    t = doc.add_table(rows=1, cols=2); no_borders(t)
    t.columns[0].width = Cm(2.2); t.columns[1].width = Cm(15.8)
    c0 = t.cell(0,0).paragraphs[0]; c0.paragraph_format.space_after=Pt(0)
    run(c0, num, bold=True, color=LEAF, size=30)
    c1 = t.cell(0,1).paragraphs[0]; c1.paragraph_format.space_after=Pt(0)
    c1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rr = run(c1, title, bold=True, color=GREEN_800, size=20)
    t.cell(0,1).vertical_alignment = 1  # center
    # linha inferior
    p = doc.add_paragraph(); p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(6)
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom"); b.set(qn("w:val"),"single"); b.set(qn("w:sz"),"18")
    b.set(qn("w:space"),"1"); b.set(qn("w:color"),GREEN_600); pbdr.append(b); pPr.append(pbdr)

def intro(text):
    p = para(); run(p, text, color=MUTED, italic=True, size=10.5)
    p.paragraph_format.space_after = Pt(8)

def kv_table(rows, widths, header=None):
    cols = len(rows[0])
    t = doc.add_table(rows=0, cols=cols); t.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_borders(t, LINE, 4, "all")
    if header:
        hr = t.add_row().cells
        for i, h in enumerate(header):
            cell_bg(t.cell(len(t.rows)-1, i), GREEN_700)
            pp = hr[i].paragraphs[0]; pp.paragraph_format.space_after=Pt(1); pp.paragraph_format.space_before=Pt(1)
            run(pp, h, bold=True, color=WHITE, size=9)
    for ri, rowdata in enumerate(rows):
        cells = t.add_row().cells
        if ri % 2 == 1:
            for i in range(cols): cell_bg(t.cell(len(t.rows)-1,i), SOFT)
        for i, (txt, mono) in enumerate(rowdata):
            pp = cells[i].paragraphs[0]; pp.paragraph_format.space_after=Pt(1); pp.paragraph_format.space_before=Pt(1)
            run(pp, txt, mono=mono, size=9, color=(GREEN_800 if mono else INK))
    for i, w in enumerate(widths):
        for c in t.columns[i].cells: c.width = Cm(w)
    return t

def callout(title, body, kind="leaf"):
    color = {"leaf":LEAF,"warn":C_PUT,"info":C_POST}[kind]
    fill  = {"leaf":"F1F7EC","warn":"FBF4E6","info":"EEF4FA"}[kind]
    t = doc.add_table(rows=1, cols=1); no_borders(t); t.alignment=WD_TABLE_ALIGNMENT.LEFT
    t.columns[0].width = Cm(17)
    cell = t.cell(0,0); cell_bg(cell, fill); left_accent(cell, color)
    if title:
        pt = cell.paragraphs[0]; pt.paragraph_format.space_after=Pt(1)
        run(pt, title, bold=True, color=color if kind!="leaf" else GREEN_800, size=10.5)
        pb = cell.add_paragraph()
    else:
        pb = cell.paragraphs[0]
    run(pb, body, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def endpoint(verb, path, access, desc, extra=None):
    cmap = {"GET":C_GET,"POST":C_POST,"PUT":C_PUT,"DELETE":C_DEL}
    amap = {"Público":MUTED,"Autenticado":GREEN_600,"Admin":C_DEL}
    t = doc.add_table(rows=1, cols=3); set_borders(t, LINE, 4, "box")
    t.columns[0].width = Cm(2.2); t.columns[1].width = Cm(11.3); t.columns[2].width = Cm(3.5)
    # verb
    cv = t.cell(0,0); cell_bg(cv, cmap[verb])
    pv = cv.paragraphs[0]; pv.alignment=WD_ALIGN_PARAGRAPH.CENTER
    pv.paragraph_format.space_after=Pt(1); pv.paragraph_format.space_before=Pt(1)
    run(pv, verb, bold=True, color=WHITE, size=8.5, mono=True)
    cv.vertical_alignment = 1
    # path
    cp = t.cell(0,1); cell_bg(cp, SOFT)
    pp = cp.paragraphs[0]; pp.paragraph_format.space_after=Pt(1); pp.paragraph_format.space_before=Pt(1)
    run(pp, path, bold=True, mono=True, size=10, color=INK); cp.vertical_alignment=1
    # access
    ca = t.cell(0,2); cell_bg(ca, SOFT)
    pa = ca.paragraphs[0]; pa.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    pa.paragraph_format.space_after=Pt(1); pa.paragraph_format.space_before=Pt(1)
    run(pa, access.upper(), bold=True, color=amap[access], size=7.5); ca.vertical_alignment=1
    # descrição
    pd = doc.add_paragraph(); pd.paragraph_format.left_indent=Cm(0.2)
    pd.paragraph_format.space_before=Pt(2); pd.paragraph_format.space_after=Pt(2)
    run(pd, desc, size=9.5)
    if extra:
        for lbl, val in extra:
            pe = doc.add_paragraph(); pe.paragraph_format.left_indent=Cm(0.2); pe.paragraph_format.space_after=Pt(1)
            run(pe, lbl+"  ", bold=True, color=MUTED, size=8)
            run(pe, val, mono=True, size=8.5, color=GREEN_800)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)

def card_grid(cards):
    t = doc.add_table(rows=0, cols=2); no_borders(t)
    for i in range(0, len(cards), 2):
        cells = t.add_row().cells
        for j in range(2):
            if i+j >= len(cards): continue
            tag, title, body = cards[i+j]
            cell = cells[j]; cell_bg(cell, SOFT)
            p1 = cell.paragraphs[0]; p1.paragraph_format.space_after=Pt(0)
            run(p1, tag.upper(), bold=True, color=GREEN_600, size=7.5)
            p2 = cell.add_paragraph(); p2.paragraph_format.space_after=Pt(1)
            run(p2, title, bold=True, color=GREEN_800, size=11)
            p3 = cell.add_paragraph(); p3.paragraph_format.space_after=Pt(2)
            run(p3, body, size=9.5)
        for c in t.columns[0].cells: c.width = Cm(8.3)
        for c in t.columns[1].cells: c.width = Cm(8.3)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

# =================================================================
# CAPA — seção sem margens com imagem em página cheia
# =================================================================
sec0 = doc.sections[0]
sec0.page_width = Cm(21); sec0.page_height = Cm(29.7)
sec0.left_margin = sec0.right_margin = sec0.top_margin = sec0.bottom_margin = Cm(0)
pcover = doc.paragraphs[0]; pcover.paragraph_format.space_after = Pt(0)
pcover.alignment = WD_ALIGN_PARAGRAPH.CENTER
pcover.add_run().add_picture("docs/_cover.png", width=Cm(21))

# =================================================================
# SEÇÃO DE CONTEÚDO — margens normais + cabeçalho/rodapé
# =================================================================
sec = doc.add_section(WD_SECTION.NEW_PAGE)
sec.page_width = Cm(21); sec.page_height = Cm(29.7)
sec.left_margin = sec.right_margin = Cm(2)
sec.top_margin = Cm(1.6); sec.bottom_margin = Cm(1.4)
sec.header_distance = Cm(0.9); sec.footer_distance = Cm(0.8)
sec.header.is_linked_to_previous = False
sec.footer.is_linked_to_previous = False

# header
hp = sec.header.paragraphs[0]
htab = hp.paragraph_format
hp.text = ""
run(hp, "GRUPO FACHOLI · SEMENTES E NUTRIÇÃO", bold=True, color=GREEN_700, size=7.5)
tabs = hp.paragraph_format.tab_stops
from docx.enum.text import WD_TAB_ALIGNMENT
tabs.add_tab_stop(Cm(17), WD_TAB_ALIGNMENT.RIGHT)
r = hp.add_run("\tDOCUMENTAÇÃO TÉCNICA — API DE ORIGINAÇÃO")
r.bold = True; r.font.size = Pt(7.5); r.font.color.rgb = RGBColor.from_string(MUTED)
pPr = hp._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr")
bb = OxmlElement("w:bottom"); bb.set(qn("w:val"),"single"); bb.set(qn("w:sz"),"8")
bb.set(qn("w:space"),"2"); bb.set(qn("w:color"),LEAF); pbdr.append(bb); pPr.append(pbdr)

# footer com número de página
fp = sec.footer.paragraphs[0]; fp.text=""
run(fp, "API de Originação de Gado · v1", color=MUTED, size=7.5)
ftabs = fp.paragraph_format.tab_stops
ftabs.add_tab_stop(Cm(17), WD_TAB_ALIGNMENT.RIGHT)
fr = fp.add_run("\tPágina "); fr.font.size=Pt(7.5); fr.font.color.rgb=RGBColor.from_string(MUTED)
def add_field(p, instr):
    fb = OxmlElement("w:fldSimple"); fb.set(qn("w:instr"), instr)
    rr = OxmlElement("w:r"); rpr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"),"15"); rpr.append(sz)
    col = OxmlElement("w:color"); col.set(qn("w:val"),MUTED); rpr.append(col)
    rr.append(rpr); fb.append(rr); p._p.append(fb)
add_field(fp, "PAGE")
fr2 = fp.add_run(" de "); fr2.font.size=Pt(7.5); fr2.font.color.rgb=RGBColor.from_string(MUTED)
add_field(fp, "NUMPAGES")
pPrf = fp._p.get_or_add_pPr(); pbdrf = OxmlElement("w:pBdr")
tb = OxmlElement("w:top"); tb.set(qn("w:val"),"single"); tb.set(qn("w:sz"),"6")
tb.set(qn("w:space"),"2"); tb.set(qn("w:color"),LINE); pbdrf.append(tb); pPrf.append(pbdrf)

# =================================================================
# SUMÁRIO
# =================================================================
heading_p = doc.add_paragraph(); heading_p.paragraph_format.space_after=Pt(4)
run(heading_p, "Sumário", bold=True, color=GREEN_800, size=22)
p = doc.add_paragraph(); pPr=p._p.get_or_add_pPr(); pbdr=OxmlElement("w:pBdr")
bb=OxmlElement("w:bottom"); bb.set(qn("w:val"),"single"); bb.set(qn("w:sz"),"18"); bb.set(qn("w:space"),"1"); bb.set(qn("w:color"),GREEN_600); pbdr.append(bb); pPr.append(pbdr)
intro("Navegue pelas seções para localizar rapidamente as informações técnicas da API.")
toc = [("01","Visão Geral","Propósito e capacidades"),
       ("02","Arquitetura e Tecnologias","Stack e camadas"),
       ("03","Autenticação e Autorização","JWT e perfis"),
       ("04","Convenções da API","Erros, paginação, unidades"),
       ("05","Modelo de Dados","Entidades e campos"),
       ("06","Referência de Endpoints","Todos os recursos"),
       ("07","Regras de Cálculo e Negócio","Fórmulas e validações"),
       ("08","Execução e Configuração","Como rodar e Swagger"),
       ("09","Suporte e Considerações Finais","Contato")]
tt = doc.add_table(rows=0, cols=3); set_borders(tt, LINE, 4, "insideH_only" if False else "all")
no_borders(tt)
for num,name,desc in toc:
    cells = tt.add_row().cells
    p0=cells[0].paragraphs[0]; run(p0,num,bold=True,color=LEAF,size=13)
    p1=cells[1].paragraphs[0]; run(p1,name,bold=True,color=GREEN_800,size=11.5)
    p2=cells[2].paragraphs[0]; p2.alignment=WD_ALIGN_PARAGRAPH.RIGHT; run(p2,desc,color=MUTED,size=9.5)
    # linha divisória inferior em cada célula
    for c in cells:
        tcPr=c._tc.get_or_add_tcPr(); bdr=OxmlElement("w:tcBorders")
        e=OxmlElement("w:bottom"); e.set(qn("w:val"),"single"); e.set(qn("w:sz"),"4"); e.set(qn("w:space"),"0"); e.set(qn("w:color"),LINE); bdr.append(e); tcPr.append(bdr)
tt.columns[0].width=Cm(1.6); tt.columns[1].width=Cm(8.4); tt.columns[2].width=Cm(7)

# =================================================================
# 01 VISÃO GERAL
# =================================================================
section_header("01","Visão Geral")
intro("O que é a API e quais problemas ela resolve.")
para("A API de Originação de Gado (projeto PrecoBoi.Api) é o serviço de back-end da plataforma "
     "Facholi para apoio à compra e originação de bovinos. Centraliza cadastros de referência, o motor de "
     "simulação de preços de praça, a gestão do ciclo de vida das negociações e os dashboards analíticos "
     "consumidos pelo aplicativo web (PWA).")
heading("Principais capacidades", 1)
card_grid([
    ("Cadastros","Dados de referência","Corretores, municípios de origem/destino, categorias, ICMS por UF, cotações regionais e configuração de comissão — com trilha de auditoria."),
    ("Simulação","Preço de praça","Cálculo do preço na praça a partir do preço colocado (descontando frete, ICMS e comissão) e rankings de oportunidades por origem."),
    ("Negociações","Ciclo de compra","Criação, edição, fechamento e controle de entregas das negociações, com cálculo automático do preço colocado por item."),
    ("Dashboards","Indicadores","Consolidação de volumes e preços médios ponderados por comprador, corretor e categoria."),
])
callout("Glossário rápido",
    "Praça: preço de referência na região de origem.  •  Colocado: preço posto na fazenda (R$/kg).  "
    "•  Arroba (@): 30 kg.  •  Ágio/Deságio: diferença percentual frente à cotação da praça.", "info")

# =================================================================
# 02 ARQUITETURA
# =================================================================
section_header("02","Arquitetura e Tecnologias")
intro("Stack tecnológico e organização em camadas.")
heading("Stack", 1)
kv_table([
    [("Plataforma / runtime",False),(".NET 8 — ASP.NET Core Web API",False)],
    [("Acesso a dados",False),("Dapper (micro-ORM) sobre Npgsql",False)],
    [("Banco de dados",False),("PostgreSQL (snake_case → PascalCase)",False)],
    [("Autenticação",False),("JWT Bearer (HMAC-SHA256)",False)],
    [("Hash de senha",False),("BCrypt",False)],
    [("Documentação",False),("OpenAPI 3 via Swashbuckle (Swagger UI)",False)],
], [6.4,10.6], header=["Componente","Tecnologia"])
heading("Organização em camadas", 1)
kv_table([
    [("Controllers",False),("Recebem requisições HTTP, validam o contexto do usuário e delegam para serviços/repositórios.",False)],
    [("Services",False),("Regras de negócio e cálculos: AuthService, CalculoService, NegociacaoService.",False)],
    [("Repositories",False),("Acesso a dados com Dapper; uma classe por agregado.",False)],
    [("Models",False),("Entidades de domínio persistidas.",False)],
    [("DTOs",False),("Contratos de entrada/saída expostos pela API.",False)],
], [4.5,12.5], header=["Camada","Responsabilidade"])
callout("CORS","A API libera origens configuradas para o front-end (PWA), permitindo qualquer cabeçalho/método e credenciais. As origens permitidas são definidas por configuração de ambiente.","leaf")

# =================================================================
# 03 AUTENTICAÇÃO
# =================================================================
section_header("03","Autenticação e Autorização")
intro("Como obter e usar o token de acesso.")
heading("Fluxo de autenticação", 1)
for i,t in enumerate([
    "O cliente envia e-mail e senha para POST /api/auth/login.",
    "A API valida as credenciais (senha conferida via BCrypt) e o status ativo do usuário.",
    "Em caso de sucesso, retorna um token JWT assinado, além de nome, e-mail, perfil e id.",
    "O cliente envia o token nas requisições seguintes no cabeçalho Authorization: Bearer <token>."], 1):
    p=doc.add_paragraph(style="List Number"); run(p,t,size=10.5)
callout("Claims do token","O JWT carrega NameIdentifier (id), Name (nome), Email e Role (perfil). O perfil determina o acesso às operações administrativas.","info")
heading("Perfis de acesso", 1)
card_grid([
    ("Perfil","Administrador","Acesso total: gestão de usuários, todos os cadastros de referência, configuração de comissão e consulta à auditoria."),
    ("Perfil","Comprador","Operação de negociações, simulações e dashboards. Pode editar/excluir apenas as negociações que criou."),
])
kv_table([
    [("Login e diagnóstico de fuso horário",False),("Público",False)],
    [("Leitura de cadastros, simulação, negociações, dashboards",False),("Autenticado",False)],
    [("Escrita em cadastros, usuários e auditoria",False),("Admin",False)],
], [11,6], header=["Tipo de operação","Acesso exigido"])
callout("Respostas de autenticação","Requisições sem token válido retornam 401 Unauthorized. Usuários autenticados sem o perfil exigido recebem 403 Forbidden.","warn")

# =================================================================
# 04 CONVENÇÕES
# =================================================================
section_header("04","Convenções da API")
intro("Padrões transversais a todos os endpoints.")
heading("Base e formato", 1)
for t in ["Todas as rotas têm o prefixo /api e trafegam application/json.",
          "Datas e horas seguem o fuso do servidor (America/São_Paulo).",
          "Valores monetários em Reais (R$). Preços em R$/kg ou R$/@ (arroba = 30 kg)."]:
    p=doc.add_paragraph(style="List Bullet"); run(p,t,size=10.5)
heading("Tratamento de erros", 1)
para("Erros de negócio e validação retornam 400 Bad Request com um corpo no formato:")
kv_table([[("mensagem",True),("Texto explicativo do erro, adequado para exibição ao usuário.",False)]],
         [4,13], header=["Campo","Descrição"])
heading("Códigos de status usados", 1)
kv_table([
    [("200 OK",True),("Sucesso com corpo de resposta.",False)],
    [("201 Created",True),("Recurso criado (retorna o recurso).",False)],
    [("204 No Content",True),("Sucesso sem corpo (atualizações/exclusões).",False)],
    [("400 Bad Request",True),("Dados inválidos ou regra de negócio violada.",False)],
    [("401 Unauthorized",True),("Token ausente ou inválido.",False)],
    [("403 Forbidden",True),("Sem permissão (perfil insuficiente).",False)],
    [("404 Not Found",True),("Recurso não encontrado.",False)],
    [("409 Conflict",True),("Conflito (ex.: e-mail já cadastrado).",False)],
], [4,13], header=["Código","Significado"])
heading("Paginação", 1)
para("Listagens paginadas aceitam pagina e tamanhoPagina na query string e retornam o envelope:")
kv_table([
    [("items",True),("Coleção de registros da página.",False)],
    [("total",True),("Total de registros que atendem ao filtro.",False)],
    [("pagina",True),("Página atual.",False)],
    [("tamanhoPagina",True),("Quantidade de itens por página.",False)],
], [4,13], header=["Campo","Descrição"])

# =================================================================
# 05 MODELO DE DADOS
# =================================================================
section_header("05","Modelo de Dados")
intro("Principais entidades de domínio e seus campos.")
def entity(title, rows):
    heading(title, 2)
    kv_table(rows, [4.5,3,9.5], header=["Campo","Tipo","Descrição"])
entity("Usuário", [
    [("id",True),("int",False),("Identificador.",False)],
    [("nome",True),("string",False),("Nome do usuário.",False)],
    [("email",True),("string",False),("E-mail (usado no login).",False)],
    [("telefone",True),("string",False),("Telefone de contato.",False)],
    [("perfil",True),("string",False),("Admin ou Comprador.",False)],
    [("ativo",True),("bool",False),("Se pode acessar o sistema.",False)],
    [("criadoEm",True),("datetime",False),("Data de criação.",False)]])
entity("Corretor", [
    [("id",True),("int",False),("Identificador.",False)],
    [("nome / telefone",True),("string",False),("Identificação e contato.",False)],
    [("municipio / uf",True),("string",False),("Localização do corretor.",False)],
    [("propriedade",True),("string",False),("Propriedade associada.",False)],
    [("observacoes",True),("string",False),("Anotações livres.",False)],
    [("ativo",True),("bool",False),("Status do corretor.",False)]])
entity("Município de Origem", [
    [("nome / uf",True),("string",False),("Município e UF.",False)],
    [("distanciaKm",True),("decimal",False),("Distância até o destino (km).",False)],
    [("valorKm",True),("decimal",False),("Valor do frete por km.",False)],
    [("ativo",True),("bool",False),("Disponível para simulação/negociação.",False)]])
entity("Categoria", [
    [("nome",True),("string",False),("Ex.: Bezerro, Garrote, Boi.",False)],
    [("pesoMin / pesoMax",True),("decimal",False),("Faixa de peso (kg).",False)],
    [("pesoMedio",True),("decimal",False),("Peso médio usado nos cálculos.",False)],
    [("cabCaminhao",True),("int",False),("Cabeças por caminhão (rateio do frete).",False)],
    [("ordem",True),("int",False),("Ordem de exibição.",False)]])
entity("ICMS (por UF)", [
    [("uf",True),("string",False),("Unidade federativa.",False)],
    [("aliquota",True),("decimal",False),("Alíquota de ICMS.",False)],
    [("recuperacao",True),("decimal",False),("Percentual de recuperação.",False)],
    [("icmsEfetivo",True),("decimal",False),("ICMS efetivo aplicado no cálculo.",False)]])
entity("Cotação Regional", [
    [("uf",True),("string",False),("UF da praça.",False)],
    [("pracaReferenciaUf",True),("string?",False),("UF de referência (opcional).",False)],
    [("valorArroba",True),("decimal",False),("Cotação da arroba (R$/@).",False)],
    [("agios[]",True),("lista",False),("Ágio percentual por categoria.",False)]])
entity("Negociação & Itens", [
    [("numero",True),("string",False),("Identificador legível NNN/AAAA.",False)],
    [("compradorId / corretorId",True),("int",False),("Partes da negociação.",False)],
    [("municipioOrigemId / destinoId",True),("int",False),("Rota da operação.",False)],
    [("dataPrevistaEntrega",True),("datetime?",False),("Previsão de entrega.",False)],
    [("status",True),("string",False),("EmNegociacao ou Fechado.",False)],
    [("itens[].categoriaId",True),("int",False),("Categoria do item.",False)],
    [("itens[].qtdNegociada",True),("int?",False),("Cabeças negociadas.",False)],
    [("itens[].precoNegociado",True),("decimal?",False),("R$/kg na praça (origem).",False)],
    [("itens[].pesoMedio",True),("decimal?",False),("Peso médio do item.",False)],
    [("itens[].precoColocado",True),("decimal?",False),("R$/kg posto fazenda (calculado).",False)],
    [("itens[].qtdEntregue",True),("int",False),("Cabeças entregues.",False)],
    [("itens[].statusEntrega",True),("string",False),("Pendente / Parcial / Concluido.",False)]])

# =================================================================
# 06 ENDPOINTS
# =================================================================
section_header("06","Referência de Endpoints")
intro("Todos os recursos expostos, agrupados por área. Acesso: Público · Autenticado · Admin.")
heading("Autenticação", 1)
endpoint("POST","/api/auth/login","Público","Autentica e retorna o token JWT.",
         [("Corpo","email, senha"),("Respostas","200 token+dados · 401 inválido/inativo")])
endpoint("GET","/api/auth/me","Autenticado","Retorna os dados do usuário do token atual.",
         [("Respostas","200 usuário · 404 não encontrado")])
heading("Usuários", 1)
endpoint("GET","/api/usuarios","Autenticado","Lista usuários. Filtro opcional ?ativo=true|false.")
endpoint("GET","/api/usuarios/{id}","Autenticado","Obtém um usuário. 404 se não existir.")
endpoint("POST","/api/usuarios","Admin","Cria um usuário (senha armazenada com hash BCrypt).",
         [("Corpo","nome, email, senha, telefone, perfil, ativo"),("Respostas","201 criado · 409 e-mail já cadastrado")])
endpoint("PUT","/api/usuarios/{id}","Admin","Atualiza um usuário. A senha só muda se 'senha' for informada.")
endpoint("DELETE","/api/usuarios/{id}","Admin","Exclui um usuário.")
heading("Cadastros — Corretores", 1)
endpoint("GET","/api/corretores","Autenticado","Lista corretores. Filtro opcional ?ativo.")
endpoint("GET","/api/corretores/{id}","Autenticado","Obtém um corretor.")
endpoint("POST","/api/corretores","Admin","Cria um corretor (gera auditoria).",
         [("Corpo","nome, telefone, municipio, uf, propriedade, observacoes, ativo")])
endpoint("PUT","/api/corretores/{id}","Admin","Atualiza um corretor.")
endpoint("DELETE","/api/corretores/{id}","Admin","Exclui um corretor.")
heading("Cadastros — Municípios", 1)
endpoint("GET","/api/municipios-origem","Autenticado","Lista municípios de origem (distância e valor por km). Filtro ?ativo.")
endpoint("GET","/api/municipios-origem/{id}","Autenticado","Obtém um município de origem.")
endpoint("POST","/api/municipios-origem","Admin","Cria um município de origem.",
         [("Corpo","nome, uf, distanciaKm, valorKm, ativo")])
endpoint("PUT","/api/municipios-origem/{id}","Admin","Atualiza um município. Mudança no valor por km é auditada.")
endpoint("DELETE","/api/municipios-origem/{id}","Admin","Exclui um município de origem.")
endpoint("GET","/api/municipios-destino","Autenticado","Lista municípios de destino.")
endpoint("GET","/api/municipios-destino/padrao","Autenticado","Retorna o município de destino padrão.")
endpoint("POST","/api/municipios-destino","Admin","Cria um município de destino. Corpo: nome, uf, padrao.")
endpoint("PUT","/api/municipios-destino/{id}","Admin","Atualiza um município de destino.")
endpoint("DELETE","/api/municipios-destino/{id}","Admin","Exclui um município de destino.")
heading("Cadastros — Categorias", 1)
endpoint("GET","/api/categorias","Autenticado","Lista categorias e seus parâmetros.")
endpoint("POST","/api/categorias","Admin","Cria categoria. Valida faixa de peso, peso médio na faixa e cabeças por caminhão > 0.",
         [("Corpo","nome, pesoMin, pesoMax, pesoMedio, cabCaminhao, ordem, agioPadrao?")])
endpoint("PUT","/api/categorias/{id}","Admin","Atualiza uma categoria.")
endpoint("DELETE","/api/categorias/{id}","Admin","Exclui uma categoria. Bloqueada se estiver em uso em negociações (400).")
heading("Cadastros — ICMS, Cotações e Comissão", 1)
endpoint("GET","/api/icms","Autenticado","Lista alíquotas de ICMS e recuperação por UF.")
endpoint("PUT","/api/icms/{uf}","Admin","Atualiza o ICMS de uma UF. Corpo: aliquota, recuperacao. Mudança auditada.")
endpoint("GET","/api/cotacoes","Autenticado","Lista cotações regionais (valor da arroba e ágios por categoria).")
endpoint("GET","/api/cotacoes/{uf}","Autenticado","Obtém a cotação de uma UF. 404 se não cadastrada.")
endpoint("POST","/api/cotacoes","Admin","Cria/atualiza (upsert) a cotação de uma UF. Mudança no valor da arroba é auditada.",
         [("Corpo","uf, pracaReferenciaUf?, valorArroba, agios[] (categoriaId, percentual)")])
endpoint("GET","/api/config-comissao","Autenticado","Obtém a configuração de comissão vigente.")
endpoint("POST","/api/config-comissao","Admin","Salva a configuração de comissão. Corpo: percentual, ativo. Mudança auditada.")
heading("Auditoria", 1)
endpoint("GET","/api/auditoria","Admin","Lista o trilho de auditoria, paginado.",
         [("Filtros","tabela, usuarioId, dataInicio, dataFim, pagina, tamanhoPagina")])
heading("Simulação", 1)
endpoint("POST","/api/simulacao","Autenticado","Calcula o preço de praça por categoria para uma rota origem→destino.",
         [("Corpo","municipioOrigemId, municipioDestinoId, itens[] (categoriaId, precoColocado)")])
endpoint("GET","/api/simulacao/oportunidades","Autenticado","Ranking (Modo A — com ICMS) por categoria e preço colocado, ordenado pelo maior ágio.",
         [("Query","categoriaId, precoColocado (> 0)")])
endpoint("GET","/api/simulacao/oportunidades-praca","Autenticado","Ranking (Modo B — sem ICMS) pelo menor custo colocado. Origens sem cotação ativa são ignoradas.",
         [("Query","categoriaId")])
endpoint("GET","/api/simulacao/rapida","Autenticado","Frete por kg de todas as categorias para uma rota.",
         [("Query","origemId, destinoId")])
heading("Negociações", 1)
endpoint("GET","/api/negociacoes","Autenticado","Lista negociações paginadas.",
         [("Filtros","compradorId, corretorId, categoria, uf, cidadeOrigem, status, pagina, tamanhoPagina")])
endpoint("GET","/api/negociacoes/{id}","Autenticado","Obtém uma negociação completa (com itens).")
endpoint("POST","/api/negociacoes","Autenticado","Cria uma negociação. Exige ao menos uma categoria com quantidade > 0; preço colocado calculado automaticamente.",
         [("Corpo","compradorId, corretorId, municipioOrigemId, municipioDestinoId, dataPrevistaEntrega?, observacoes?, itens[]")])
endpoint("PUT","/api/negociacoes/{id}","Autenticado","Atualiza uma negociação. Fechada não pode ser editada; comprador só edita as próprias (403); alterações auditadas item a item.")
endpoint("DELETE","/api/negociacoes/{id}","Autenticado","Exclui uma negociação. Comprador só exclui as próprias (403).")
endpoint("POST","/api/negociacoes/{id}/fechar","Autenticado","Fecha a negociação (status → Fechado). Já fechada retorna 400.")
endpoint("PUT","/api/negociacoes/entrega","Autenticado","Registra quantidades entregues por item. Apenas para negociações fechadas; atualiza o status de entrega.",
         [("Corpo","negociacaoId, itens[] (itemId, qtdEntregue)")])
heading("Dashboard", 1)
endpoint("GET","/api/dashboard/totais","Autenticado","Totais consolidados (quantidade e preços médios ponderados) da base filtrada.")
endpoint("GET","/api/dashboard/compradores","Autenticado","Indicadores agregados por comprador.")
endpoint("GET","/api/dashboard/compradores/{id}/negociacoes","Autenticado","Negociações de um comprador específico.")
endpoint("GET","/api/dashboard/compradores/{id}/categorias-corretor","Autenticado","Distribuição por corretor e categoria das negociações do comprador.")
endpoint("GET","/api/dashboard/corretores","Autenticado","Indicadores agregados por corretor e categoria.")
endpoint("GET","/api/dashboard/por-categoria","Autenticado","Indicadores agregados por categoria.")
endpoint("GET","/api/dashboard/por-categoria/{id}/detalhe","Autenticado","Quebra de uma categoria por comprador e corretor.")
endpoint("GET","/api/dashboard/resumo-cabecas","Autenticado","Resumo de cabeças em andamento e fechadas, por categoria e total geral.")
heading("Diagnóstico", 1)
endpoint("GET","/api/debug/tz","Público","Compara fuso e horário do servidor (.NET) e do PostgreSQL.")

# =================================================================
# 07 REGRAS DE CÁLCULO
# =================================================================
section_header("07","Regras de Cálculo e Negócio")
intro("Fórmulas do motor de precificação e validações de domínio.")
heading("Frete por quilograma", 1)
para("Base de todos os cálculos. Considera ida e volta (×2) e um fator de recuperação de 88%:")
callout("", "FreteKg = (DistânciaKm × 2 × ValorKm × 0,88) ÷ CabCaminhão ÷ PesoMédio", "leaf")
heading("Direção 1 — Colocado → Praça (Simulação)", 1)
para("A partir do preço posto na fazenda, deduz frete, ICMS e comissão para chegar ao preço de praça:")
kv_table([
    [("Valor na compra",False),("PrecoColocado − FreteKg",True)],
    [("ICMS",False),("ValorNaCompra × IcmsEfetivo",True)],
    [("Comissão",False),("ValorNaCompra × (Percentual ÷ 100) — se ativa",True)],
    [("Preço de praça",False),("ValorNaCompra − ICMS − Comissão (arred. p/ baixo, 1 casa)",True)],
], [6,11], header=["Passo","Fórmula"])
heading("Direção 2 — Negociado → Colocado (Negociações)", 1)
para("Cálculo inverso usado ao registrar itens. ICMS e comissão incidem sobre o preço negociado:")
callout("", "PrecoColocado = PrecoNegociado + FreteKg + (PrecoNegociado × IcmsEfetivo) + (PrecoNegociado × Comissão%)", "leaf")
heading("Direção 3 — Cotação da praça (Modo B)", 1)
para("Custo de buscar o animal na origem, sem ICMS:")
kv_table([
    [("Cotação da praça (R$/kg)",False),("(ValorArroba ÷ 30) × (1 + Ágio% da categoria)",True)],
    [("Custo colocado (R$/kg)",False),("CotacaoPracaKg + FreteKg",True)],
], [6,11], header=["Passo","Fórmula"])
heading("Deságio (Modo A)", 1)
para("Compara o preço de praça calculado com a cotação crua da praça (sem ágio):")
callout("", "Deságio% = (PrecoPraca ÷ (ValorArroba ÷ 30) − 1) × 100  ·  Positivo = acima da praça; Negativo = abaixo.", "leaf")
heading("Validações de negociação", 1)
for t in ["Pelo menos uma categoria deve ter quantidade de cabeças maior que zero.",
          "Observações são limitadas a 500 caracteres.",
          "Negociação fechada não pode ser editada.",
          "Comprador só pode editar/excluir negociações que criou; Admin não tem essa restrição.",
          "Controle de entrega só para negociações fechadas; qtdEntregue não pode ser negativa.",
          "Status de entrega: Pendente (0), Concluido (≥ negociada) ou Parcial."]:
    p=doc.add_paragraph(style="List Bullet"); run(p,t,size=10.5)
callout("Auditoria automática","Criação, edição, exclusão, fechamento e entrega de negociações — bem como alterações de parâmetros (valor por km, ICMS, cotação, comissão) — são registradas na trilha de auditoria com usuário, valores anterior/novo e descrição.","info")

# =================================================================
# 08 EXECUÇÃO
# =================================================================
section_header("08","Execução e Configuração")
intro("Como executar a API localmente e acessar a documentação interativa.")
heading("Pré-requisitos", 1)
for t in [".NET SDK 8.0","Acesso a uma instância PostgreSQL"]:
    p=doc.add_paragraph(style="List Bullet"); run(p,t,size=10.5)
heading("Executando", 1)
kv_table([
    [("dotnet restore",True),("Restaura os pacotes.",False)],
    [("dotnet build",True),("Compila o projeto.",False)],
    [("dotnet run",True),("Sobe a API.",False)],
], [6,11], header=["Comando","Descrição"])
heading("Documentação interativa (Swagger)", 1)
para("Com a API em execução, a documentação OpenAPI fica disponível na rota /swagger. A interface inclui o "
     "botão Authorize para informar o token JWT e testar os endpoints protegidos direto do navegador.")
heading("Configuração por ambiente", 1)
para("As configurações sensíveis são fornecidas por variáveis de ambiente / arquivo de configuração do ambiente, "
     "não devendo ser versionadas com valores reais:")
kv_table([
    [("ConnectionStrings:DefaultConnection",True),("String de conexão com o PostgreSQL.",False)],
    [("Jwt:Key",True),("Chave secreta de assinatura do JWT.",False)],
    [("Jwt:Issuer / Jwt:Audience",True),("Emissor e público do token.",False)],
    [("AllowedOrigins",True),("Origens liberadas no CORS (separadas por vírgula).",False)],
], [7,10], header=["Chave","Descrição"])
callout("Boas práticas de segurança","Mantenha string de conexão, credenciais e chave JWT fora do controle de versão (variáveis de ambiente ou gerenciador de segredos). Rotacione qualquer segredo exposto. Os valores reais foram omitidos desta documentação intencionalmente.","warn")

# =================================================================
# 09 SUPORTE
# =================================================================
section_header("09","Suporte e Considerações Finais")
intro("Canais de apoio e encerramento.")
para("Esta documentação descreve a superfície técnica da API de Originação na versão v1. Para a documentação "
     "sempre atualizada e testável, consulte o Swagger da própria aplicação em /swagger, que reflete o código em execução.")
card_grid([
    ("Documentação viva","Swagger / OpenAPI","Disponível em /swagger com todos os endpoints, esquemas e autenticação JWT integrada."),
    ("Suporte","Equipe técnica","Em caso de dúvidas sobre integração ou comportamento da API, acione a equipe de desenvolvimento responsável."),
])
callout("Versionamento","Alterações de contrato devem ser refletidas tanto no código (comentários XML / anotações) quanto nesta documentação, mantendo ambas as fontes consistentes.","leaf")
pf=doc.add_paragraph(); pf.alignment=WD_ALIGN_PARAGRAPH.CENTER; pf.paragraph_format.space_before=Pt(18)
run(pf,"© 2026 Grupo Facholi — Sementes e Nutrição · Documentação Técnica da API de Originação de Gado",color=MUTED,size=9)

# =================================================================
# METADADOS NEUTROS
# =================================================================
cp = doc.core_properties
cp.author = "Grupo Facholi"
cp.last_modified_by = "Grupo Facholi"
cp.title = "Documentação Técnica — API de Originação de Gado"
cp.subject = "Referência técnica da API"
cp.category = "Documentação Técnica"
cp.comments = ""
cp.keywords = ""

doc.save("docs/Documentacao_Tecnica_API_Facholi.docx")
print("DOCX salvo.")
